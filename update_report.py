from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document('/root/.openclaw/workspace/litmatch-tracker/Litmatch海外重点国家信息报告.docx')

# 找到 Google Play 部分的起止索引
gp_start = None
gp_end = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text == '5.1 Google Play':
        gp_start = i
    if gp_start and text == '5.1 App Store':
        gp_end = i
        break

print(f"Google Play section: {gp_start} to {gp_end}")

# 删除旧内容（清空段落）
for i in range(gp_start + 1, gp_end):
    doc.paragraphs[i].text = ''

# 在 5.1 Google Play 标题后插入新内容
# 获取标题段落的参考位置
title_para = doc.paragraphs[gp_start]

# 新内容（校准后的Google Play政策）
new_content = [
    ("▸ 联系人权限：强制迁移至 Android Contact Picker [2026-04-15] 【极高】", "Normal"),
    ("Google Play 于 2026 年 4 月 15 日发布新政策：非核心功能应用必须改用 Android Contact Picker（系统级联系人选择器）获取联系人，禁止请求完整 READ_CONTACTS 权限。社交/约会类应用的「批量导入通讯录找好友」功能将受到重大限制——每次需用户手动选择单个联系人。确需完整通讯录的应用（如拨号器、通讯录管理工具）可提交 Play Developer Declaration 申请保留，但社交类应用获批概率很低。", "Normal"),
    ("执行梯度：2026-05-15 首波 enforcement（警告 → Listing suppression → 移除）；2026-10-27 Play Console 自动拦截不合规应用。", "Normal"),
    ("来源：Google 官方开发者博客 (2026-04-15) | https://android-developers.googleblog.com/2026/04/giving-users-clearer-choice-and-everyone-a-safer-more-trusted-app-ecosystem.html", "Normal"),
    ("来源：ASOWorld 政策分析 (2026-04-15) | https://asoworld.com/blog/april-2026-google-play-policy-updates/", "Normal"),
    ("", "Normal"),
    
    ("▸ 位置权限：强制采用 Location Button 模式 [2026-04-15] 【极高】", "Normal"),
    ("Google Play 更新位置权限政策：一次性精确位置访问必须采用系统级「位置按钮」（Location Button）模式，用户每次主动点击触发一次性授权，应用仅获取该次位置坐标。后台持续定位（ACCESS_BACKGROUND_LOCATION）不再作为默认选项，确需持续定位的应用（如导航、健身追踪）须提交 Play Developer Declaration 说明理由。", "Normal"),
    ("对社交应用影响：「附近的人」实时刷新列表将变得困难——每次刷新需用户手动触发位置按钮。Geofencing 不再是 foreground services 的批准用例，必须迁移至专用 Geofence API。", "Normal"),
    ("执行梯度：2026-05-15 首波 enforcement；2026-10-27 自动拦截。", "Normal"),
    ("来源：Google 官方开发者博客 (2026-04-15) | https://android-developers.googleblog.com/2026/04/giving-users-clearer-choice-and-everyone-a-safer-more-trusted-app-ecosystem.html", "Normal"),
    ("来源：HelpNetSecurity (2026-04-16) | https://www.helpnetsecurity.com/2026/04/16/google-play-store-policy-updates/", "Normal"),
    ("", "Normal"),
    
    ("▸ 账户转移：强制官方流程 + 7天安全延迟 [2026-04-15] 【高】", "Normal"),
    ("Google Play 正式规定：开发者账号所有权转移必须通过 Play Console 官方「Transfer ownership」流程，包含强制 7 天安全延迟期（冷却期），期间任一方可取消。该政策主要打击私下买卖开发者账号（credential sharing）的黑产行为。正常公司内部账号迁移不受影响，但公司并购/分拆需提前至少 7 天规划。", "Normal"),
    ("执行截止：2026-05-27。", "Normal"),
    ("来源：ASOWorld (2026-04-15) | https://asoworld.com/blog/april-2026-google-play-policy-updates/", "Normal"),
    ("来源：HelpNetSecurity (2026-04-16) | https://www.helpnetsecurity.com/2026/04/16/google-play-store-policy-updates/", "Normal"),
    ("", "Normal"),
    
    ("▸ 预测市场试点注册 [2026-04-15] 【中】", "Normal"),
    ("Google Play 推出预测市场应用全球试点计划，允许真实货币交易的预测市场应用上架，但必须于 2026-06-01 前注册加入试点计划。仅限特定地区（如美国部分州），需配合地理围栏、年龄限制和 AML 合规。社交应用如有「预测/竞猜」类功能（哪怕是虚拟积分）需评估是否触碰此红线。", "Normal"),
    ("来源：ASOWorld (2026-04-15) | https://asoworld.com/blog/april-2026-google-play-policy-updates/", "Normal"),
    ("", "Normal"),
    
    ("▸ 年龄受限内容政策澄清 [2026-04-15] 【中】", "Normal"),
    ("Google Play 澄清：约会/匹配功能仅为「附带功能」（incidental）的社交应用，不再强制要求在 Play Console 中实施「Restrict Minor Access」，但前提是应用必须维持有效的替代年龄验证机制。此条为 clarification 而非新 enforcement，Google 表示执行标准未变。", "Normal"),
    ("来源：ASOWorld (2026-04-15) | https://asoworld.com/blog/april-2026-google-play-policy-updates/", "Normal"),
    ("", "Normal"),
    
    ("▸ 约会广告政策更新 [2025-03-04] 【高】", "Normal"),
    ("Google 自 2025 年 3 月 4 日起实施新的约会与陪伴广告政策。所有约会/陪伴类广告主必须获得 Google Ads 认证才能投放广告。一般约会广告需获得「一般约会认证」；涉及敏感内容的需获得「敏感约会认证」。政策旨在提高约会广告生态的透明度和用户安全。", "Normal"),
    ("来源：We Love Digital Marketing | https://www.welovedigitalmarketing.com/blogpost/googles-dating-ads-policy/", "Normal"),
]

# 在标题后插入新段落
# 由于 python-docx 不支持在特定位置插入段落，我们采用替代方案：
# 修改现有空段落，或在末尾追加后重新排序

# 方案：清空旧段落后，用新增段落填充
# 先计算需要多少段落
import copy

# 清空 gp_start+1 到 gp_end-1 的段落
for i in range(gp_start + 1, gp_end):
    doc.paragraphs[i].text = ''

# 在 gp_start 后添加新段落（python-docx 会在文档末尾添加，所以我们需要 workaround）
# 实际上，更简单的方法：重新构建整个段落列表

# 获取 body element
body = doc.element.body

# 找到标题段落的 XML 元素
title_element = doc.paragraphs[gp_start]._element

# 在标题后逐个插入新段落
for text, style_name in new_content:
    new_p = doc.add_paragraph(text, style=style_name)
    # 移动到这个位置
    body.insert(body.index(title_element) + 1, new_p._element)
    title_element = new_p._element

# 保存
output_path = '/root/.openclaw/workspace/litmatch-tracker/Litmatch海外重点国家信息报告_v2.docx'
doc.save(output_path)
print(f"Saved to: {output_path}")
PYEOF
