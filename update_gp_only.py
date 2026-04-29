from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 创建新文档，只包含校准后的 Google Play 第8条
doc = Document()

# 标题
title = doc.add_heading('Google Play 2026 年 4 月政策更新 — 精确版', level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# 元信息
meta = doc.add_paragraph()
meta.add_run('验证状态：已交叉验证（Google 官方开发者博客 2026-04-15 + ASOWorld + HelpNetSecurity）\n').bold = True
meta.add_run('生成日期：2026-04-26 | 数据分级：已验证 ✅')

# 分隔线
doc.add_paragraph('—' * 40)

# ===== 1. 联系人权限 =====
doc.add_heading('🔴 1. 联系人权限：强制迁移至 Android Contact Picker', level=1)
p = doc.add_paragraph()
p.add_run('截止日期：').bold = True
p.add_run('2026-05-15（首波 enforcement）\n')
p.add_run('政策核心：').bold = True
p.add_run('非核心功能应用必须改用 Android Contact Picker（系统级联系人选择器），禁止请求完整 READ_CONTACTS 权限。用户每次主动选择具体联系人，应用只能获取被选中的那一条。\n')
p.add_run('对社交/约会应用的影响：').bold = True
p.add_run('「批量导入通讯录找好友」功能基本不可用。每次需用户手动选择单个联系人；如需完整通讯录匹配，可提交 Play Developer Declaration 申请保留，但社交类应用获批概率很低。\n')
p.add_run('执行梯度（渐进式，非一刀切）：').bold = True
p.add_run('警告通知 → Listing suppression（搜索降权/隐藏） → 最终移除。2026-10-27 起 Play Console 提交前自动拦截不合规应用。')

src = doc.add_paragraph()
src.add_run('来源：').bold = True
src.add_run('Google 官方开发者博客 (2026-04-15) | https://android-developers.googleblog.com/2026/04/giving-users-clearer-choice-and-everyone-a-safer-more-trusted-app-ecosystem.html\n')
src.add_run('来源：').bold = True
src.add_run('ASOWorld 政策分析 (2026-04-15) | https://asoworld.com/blog/april-2026-google-play-policy-updates/')

doc.add_paragraph()

# ===== 2. 位置权限 =====
doc.add_heading('🔴 2. 位置权限：强制采用 Location Button', level=1)
p = doc.add_paragraph()
p.add_run('截止日期：').bold = True
p.add_run('2026-05-15（首波 enforcement）\n')
p.add_run('政策核心：').bold = True
p.add_run('一次性精确位置访问必须采用系统级「位置按钮」（Location Button）模式，用户每次主动点击触发一次性授权，应用仅获取该次位置坐标。后台持续定位（ACCESS_BACKGROUND_LOCATION）不再是默认选项，确需持续定位的应用（如导航、健身追踪）须提交 Play Developer Declaration 说明理由。\n')
p.add_run('对社交/约会应用的影响：').bold = True
p.add_run('「附近的人」实时刷新列表将变得困难——每次刷新需用户手动触发位置按钮。Geofencing 不再是 foreground services 的批准用例，必须迁移至专用 Geofence API。\n')
p.add_run('执行梯度：').bold = True
p.add_run('同上，渐进式执法。2026-10-27 自动拦截。')

src = doc.add_paragraph()
src.add_run('来源：').bold = True
src.add_run('Google 官方开发者博客 (2026-04-15) | https://android-developers.googleblog.com/2026/04/giving-users-clearer-choice-and-everyone-a-safer-more-trusted-app-ecosystem.html\n')
src.add_run('来源：').bold = True
src.add_run('HelpNetSecurity (2026-04-16) | https://www.helpnetsecurity.com/2026/04/16/google-play-store-policy-updates/')

doc.add_paragraph()

# ===== 3. 账户转移 =====
doc.add_heading('🟡 3. 账户转移：强制官方流程 + 7天安全延迟', level=1)
p = doc.add_paragraph()
p.add_run('截止日期：').bold = True
p.add_run('2026-05-27\n')
p.add_run('政策核心：').bold = True
p.add_run('开发者账号所有权转移必须通过 Play Console 官方「Transfer ownership」流程，包含强制 7 天安全延迟期（冷却期），期间任一方可取消。\n')
p.add_run('对正常业务的影响：').bold = True
p.add_run('公司并购、团队重组、账号归属调整需提前至少 7 天规划。正常公司内部账号迁移不受影响。\n')
p.add_run('打击对象：').bold = True
p.add_run('私下买卖开发者账号（credential sharing）、账号倒卖黑产。')

src = doc.add_paragraph()
src.add_run('来源：').bold = True
src.add_run('ASOWorld (2026-04-15) | https://asoworld.com/blog/april-2026-google-play-policy-updates/\n')
src.add_run('来源：').bold = True
src.add_run('HelpNetSecurity (2026-04-16) | https://www.helpnetsecurity.com/2026/04/16/google-play-store-policy-updates/')

doc.add_paragraph()

# ===== 4. 预测市场 =====
doc.add_heading('🟡 4. 预测市场试点注册', level=1)
p = doc.add_paragraph()
p.add_run('截止日期：').bold = True
p.add_run('2026-06-01\n')
p.add_run('政策核心：').bold = True
p.add_run('允许真实货币交易的预测市场应用上架，但必须加入 Google 试点计划。仅限特定地区（如美国部分州），需配合地理围栏、年龄限制和 AML 合规。\n')
p.add_run('对社交/约会应用的间接警示：').bold = True
p.add_run('任何涉及「预测/竞猜/wagering」的功能（哪怕是虚拟积分）都需评估是否触碰此红线。')

src = doc.add_paragraph()
src.add_run('来源：').bold = True
src.add_run('ASOWorld (2026-04-15) | https://asoworld.com/blog/april-2026-google-play-policy-updates/')

doc.add_paragraph()

# ===== 5. 年龄受限内容澄清 =====
doc.add_heading('🟢 5. 年龄受限内容政策澄清', level=1)
p = doc.add_paragraph()
p.add_run('性质：').bold = True
p.add_run('此为 clarification，非新 enforcement。Google 表示执行标准未变，仅澄清边界。\n')
p.add_run('内容：').bold = True
p.add_run('约会/匹配功能仅为「附带功能」（incidental）的社交应用，不再强制要求在 Play Console 中实施「Restrict Minor Access」，但前提是应用必须维持有效的替代年龄验证机制。\n')
p.add_run('对 Litmatch 的影响：').bold = True
p.add_run('如果定位为「灵魂社交」而非纯约会应用，可能适用此放宽条款，但仍需有自己的年龄验证方案。')

src = doc.add_paragraph()
src.add_run('来源：').bold = True
src.add_run('ASOWorld (2026-04-15) | https://asoworld.com/blog/april-2026-google-play-policy-updates/')

doc.add_paragraph()

# ===== 6. 约会广告政策 =====
doc.add_heading('🟢 6. 约会广告政策更新', level=1)
p = doc.add_paragraph()
p.add_run('生效日期：').bold = True
p.add_run('2025-03-04\n')
p.add_run('内容：').bold = True
p.add_run('Google 实施新的约会与陪伴广告政策。所有约会/陪伴类广告主必须获得 Google Ads 认证才能投放广告。一般约会广告需获得「一般约会认证」；涉及敏感内容的需获得「敏感约会认证」。')

src = doc.add_paragraph()
src.add_run('来源：').bold = True
src.add_run('We Love Digital Marketing | https://www.welovedigitalmarketing.com/blogpost/googles-dating-ads-policy/')

doc.add_paragraph()
doc.add_paragraph('—' * 40)

# 诚实声明
warn = doc.add_heading('⚠️ 诚实声明', level=1)
p = doc.add_paragraph()
p.add_run('1. "5月15日下架"是过度简化的说法。').bold = True
p.add_run(' 实际执行是渐进的：警告 → 降权 → 移除。\n')
p.add_run('2. READ_CONTACTS 不是绝对禁止。').bold = True
p.add_run(' 可申请 Declaration 保留，但社交类获批概率很低。\n')
p.add_run('3. Location Button 不是唯一选项。').bold = True
p.add_run(' 后台定位仍可申请，需正当理由。\n')
p.add_run('4. 本文件仅基于公开信息整理，具体执行以 Google Play Policy Center 最新版本为准。')

# 保存
output_path = '/root/.openclaw/workspace/litmatch-tracker/Google_Play_2026年4月政策更新_精确版.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
PYEOF
