import json
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def load_data():
    with open('/tmp/tracker_data.json', 'r', encoding='utf-8') as f:
        return json.load(f)

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    from docx.oxml import parse_xml
    shading = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return heading

def add_paragraph_zh(doc, text, bold=False, size=10.5):
    """添加中文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.bold = bold
    return p

def format_impact(level):
    """影响级别格式化"""
    if level == '极高':
        return '【极高】'
    elif level == '高':
        return '【高】'
    elif level == '中':
        return '【中】'
    else:
        return '【低】'

def generate_report():
    data = load_data()
    
    # 统计数据
    total = len(data)
    countries = defaultdict(list)
    categories = defaultdict(list)
    high_impact = [d for d in data if d['impactLevel'] in ['高', '极高']]
    
    for item in data:
        countries[item['country']].append(item)
        categories[item['category']].append(item)
    
    # 创建文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(10.5)
    
    # ========== 封面 ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Litmatch 海外重点国家信息报告')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.color.rgb = RGBColor(31, 31, 58)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('—— 基于情报跟踪站数据（2026年4月）')
    run.font.size = Pt(14)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.color.rgb = RGBColor(100, 100, 150)
    
    doc.add_paragraph()
    
    # 报告信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(
        f'数据覆盖：{total} 条情报 | {len(countries)} 个国家/地区 | {len(categories)} 个合规领域\n'
        f'高影响级别情报：{len(high_impact)} 条 | 报告生成时间：2026-04-23'
    )
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(120, 120, 150)
    
    doc.add_page_break()
    
    # ========== 执行摘要 ==========
    add_heading_zh(doc, '一、执行摘要', level=1)
    
    summary_text = (
        f'本报告基于 Litmatch 情报跟踪站的 {total} 条公开监管情报，覆盖 {len(countries)} 个国家/地区，'
        f'涵盖 {len(categories)} 个合规领域（法律法规、政府政策、内容合规、未成年合规、反洗钱、税务合规、'
        f'执法行动、应用市场政策、同业信息）。\n\n'
        f'其中，影响级别为「高」或「极高」的情报共 {len(high_impact)} 条，'
        f'主要集中在越南（7条）、马来西亚（5条）、印度尼西亚（5条）、菲律宾（4条）、泰国（4条）等东南亚核心市场，'
        f'以及巴西、墨西哥等拉美冲刺市场。'
    )
    add_paragraph_zh(doc, summary_text)
    
    # 领域分布表
    add_heading_zh(doc, '1.1 情报领域分布', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '合规领域'
    hdr_cells[1].text = '情报数量'
    for cell in hdr_cells:
        set_cell_shading(cell, 'F0F0F5')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    for cat_name, items in sorted(categories.items(), key=lambda x: -len(x[1])):
        row_cells = table.add_row().cells
        row_cells[0].text = cat_name
        row_cells[1].text = f'{len(items)} 条'
        for cell in row_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_paragraph()
    
    # 影响级别分布
    add_heading_zh(doc, '1.2 影响级别分布', level=2)
    impact_dist = {'极高': 0, '高': 0, '中': 0, '低': 0}
    for item in data:
        impact_dist[item['impactLevel']] = impact_dist.get(item['impactLevel'], 0) + 1
    
    for level, count in impact_dist.items():
        if count > 0:
            add_paragraph_zh(doc, f'{format_impact(level)}：{count} 条')
    
    doc.add_page_break()
    
    # ========== 第二部分：东南亚核心市场 ==========
    add_heading_zh(doc, '二、东南亚核心市场', level=1)
    
    se_countries = ['菲律宾', '印度尼西亚', '泰国', '马来西亚', '越南']
    
    for country_name in se_countries:
        items = countries.get(country_name, [])
        if not items:
            continue
        
        add_heading_zh(doc, f'2.{se_countries.index(country_name)+1} {country_name}', level=2)
        
        # 按分类分组
        cat_items = defaultdict(list)
        for item in items:
            cat_items[item['category']].append(item)
        
        for cat_name in ['法律法规', '政府政策', '内容合规', '未成年合规', '税务合规', '执法行动', '反洗钱']:
            cat_data = cat_items.get(cat_name, [])
            if not cat_data:
                continue
            
            add_heading_zh(doc, f'{cat_name}（{len(cat_data)} 条）', level=3)
            
            for item in sorted(cat_data, key=lambda x: x['date'], reverse=True):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                
                # 标题
                run = p.add_run(f'▸ {item["title"]} ')
                run.bold = True
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                
                # 日期和影响
                run2 = p.add_run(f'[{item["date"]}] {format_impact(item["impactLevel"])}')
                run2.font.color.rgb = RGBColor(200, 50, 50)
                run2.font.name = 'Microsoft YaHei'
                run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                
                # 摘要
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.4)
                run3 = p2.add_run(item['summary'])
                run3.font.size = Pt(9.5)
                run3.font.color.rgb = RGBColor(80, 80, 100)
                run3.font.name = 'Microsoft YaHei'
                run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                
                # 来源
                p3 = doc.add_paragraph()
                p3.paragraph_format.left_indent = Inches(0.4)
                run4 = p3.add_run(f'来源：{item["source"]}')
                run4.font.size = Pt(9)
                run4.font.color.rgb = RGBColor(120, 120, 140)
                run4.font.italic = True
                run4.font.name = 'Microsoft YaHei'
                run4._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                
                # 关键词
                if item.get('keywords'):
                    p4 = doc.add_paragraph()
                    p4.paragraph_format.left_indent = Inches(0.4)
                    run5 = p4.add_run('关键词：' + ' | '.join(item['keywords']))
                    run5.font.size = Pt(9)
                    run5.font.color.rgb = RGBColor(100, 100, 140)
                    run5.font.name = 'Microsoft YaHei'
                    run5._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========== 第三部分：拉美冲刺市场 ==========
    add_heading_zh(doc, '三、拉美冲刺市场', level=1)
    
    la_countries = ['巴西', '墨西哥', '哥伦比亚']
    
    for country_name in la_countries:
        items = countries.get(country_name, [])
        if not items:
            continue
        
        add_heading_zh(doc, f'3.{la_countries.index(country_name)+1} {country_name}', level=2)
        
        for item in sorted(items, key=lambda x: x['date'], reverse=True):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            
            run = p.add_run(f'▸ {item["title"]} ')
            run.bold = True
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            run2 = p.add_run(f'[{item["date"]}] {format_impact(item["impactLevel"])} [{item["category"]}]')
            run2.font.color.rgb = RGBColor(200, 50, 50)
            run2.font.name = 'Microsoft YaHei'
            run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.4)
            run3 = p2.add_run(item['summary'])
            run3.font.size = Pt(9.5)
            run3.font.color.rgb = RGBColor(80, 80, 100)
            run3.font.name = 'Microsoft YaHei'
            run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_page_break()
    
    # ========== 第四部分：中东及其他市场 ==========
    add_heading_zh(doc, '四、中东及其他市场', level=1)
    
    other_countries = ['土耳其', '沙特阿拉伯', '阿联酋', '印度', '韩国', '中国台湾', '欧盟']
    
    for country_name in other_countries:
        items = countries.get(country_name, [])
        if not items:
            continue
        
        add_heading_zh(doc, f'4.{other_countries.index(country_name)+1} {country_name}', level=2)
        
        for item in sorted(items, key=lambda x: x['date'], reverse=True):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            
            run = p.add_run(f'▸ {item["title"]} ')
            run.bold = True
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            run2 = p.add_run(f'[{item["date"]}] {format_impact(item["impactLevel"])} [{item["category"]}]')
            run2.font.color.rgb = RGBColor(200, 50, 50)
            run2.font.name = 'Microsoft YaHei'
            run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.4)
            run3 = p2.add_run(item['summary'])
            run3.font.size = Pt(9.5)
            run3.font.color.rgb = RGBColor(80, 80, 100)
            run3.font.name = 'Microsoft YaHei'
            run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_page_break()
    
    # ========== 第五部分：应用市场政策 ==========
    add_heading_zh(doc, '五、应用市场政策', level=1)
    
    for market_name in ['Google Play', 'App Store']:
        items = countries.get(market_name, [])
        if not items:
            continue
        
        add_heading_zh(doc, f'5.1 {market_name}', level=2)
        
        for item in sorted(items, key=lambda x: x['date'], reverse=True):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            
            run = p.add_run(f'▸ {item["title"]} ')
            run.bold = True
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            run2 = p.add_run(f'[{item["date"]}] {format_impact(item["impactLevel"])}')
            run2.font.color.rgb = RGBColor(200, 50, 50)
            run2.font.name = 'Microsoft YaHei'
            run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.4)
            run3 = p2.add_run(item['summary'])
            run3.font.size = Pt(9.5)
            run3.font.color.rgb = RGBColor(80, 80, 100)
            run3.font.name = 'Microsoft YaHei'
            run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            if item.get('sourceUrl'):
                p3 = doc.add_paragraph()
                p3.paragraph_format.left_indent = Inches(0.4)
                run4 = p3.add_run(f'参考链接：{item["sourceUrl"]}')
                run4.font.size = Pt(9)
                run4.font.color.rgb = RGBColor(100, 100, 150)
                run4.font.italic = True
                run4.font.name = 'Microsoft YaHei'
                run4._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_page_break()
    
    # ========== 第六部分：全球趋势与同业动态 ==========
    add_heading_zh(doc, '六、全球趋势与同业动态', level=1)
    
    global_items = countries.get('全球', []) + countries.get('东南亚', []) + countries.get('拉美', [])
    
    for item in sorted(global_items, key=lambda x: x['date'], reverse=True):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        
        run = p.add_run(f'▸ {item["title"]} ')
        run.bold = True
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        run2 = p.add_run(f'[{item["date"]}] [{item["country"]}] [{item["category"]}]')
        run2.font.color.rgb = RGBColor(200, 50, 50)
        run2.font.name = 'Microsoft YaHei'
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.4)
        run3 = p2.add_run(item['summary'])
        run3.font.size = Pt(9.5)
        run3.font.color.rgb = RGBColor(80, 80, 100)
        run3.font.name = 'Microsoft YaHei'
        run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_page_break()
    
    # ========== 第七部分：合规建议与行动清单 ==========
    add_heading_zh(doc, '七、合规建议与行动清单', level=1)
    
    add_heading_zh(doc, '7.1 紧急行动项（30天内）', level=2)
    urgent_actions = [
        '【越南】确保《网络安全法》2026年7月生效前的本地实体设立和数据本地化准备',
        '【印尼】完成 PSE 注册，避免平台被封锁风险',
        '【马来西亚】评估是否需要申请 A 级 ASP 牌照，应对 MCMC 最后合规警告',
        '【Google Play】2026年5月15日前调整联系人和位置权限使用方式',
        '【税务】梳理各运营国家的 VAT/DST 注册义务，确保税务合规',
    ]
    for action in urgent_actions:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(action)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_zh(doc, '7.2 短期行动项（90天内）', level=2)
    short_actions = [
        '【未成年合规】部署年龄验证机制，适应泰国、印尼、美国多州等市场的强制要求',
        '【内容合规】建立深度伪造（Deepfake）内容检测和快速响应机制',
        '【数据保护】审查各市场数据本地化要求，更新隐私政策和同意机制',
        '【反洗钱】评估虚拟礼物/打赏系统的 AML 风险，建立交易监控机制',
    ]
    for action in short_actions:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(action)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_heading_zh(doc, '7.3 中期战略项（6个月内）', level=2)
    mid_actions = [
        '【拉美】监控巴西 CBS/IBS 税制、墨西哥代扣税等税务改革进展',
        '【中东】评估沙特内容审核要求和阿联酋广告许可制度的市场准入影响',
        '【技术】集成 Google Play 和 App Store 的年龄验证 API，适应全球趋势',
        '【竞争】对标行业安全功能（ID验证、视频验证、AI审核）提升产品竞争力',
    ]
    for action in mid_actions:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(action)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 免责声明
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        '本报告基于公开信息整理，仅供参考，不构成法律意见。\n'
        '具体合规措施请咨询当地专业律师或合规顾问。\n'
        '数据来源：Litmatch 情报跟踪站 https://summerchang2025-bot.github.io/litmatch-tracker/'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(150, 150, 170)
    run.font.italic = True
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 保存
    output_path = '/root/.openclaw/workspace/litmatch-tracker/Litmatch海外重点国家信息报告.docx'
    doc.save(output_path)
    print(f'Report saved to: {output_path}')
    
    return output_path

if __name__ == '__main__':
    generate_report()
